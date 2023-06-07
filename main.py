import pandas as pd
import streamlit as st
from st_aggrid import AgGrid
from st_aggrid import JsCode, GridUpdateMode, DataReturnMode
from st_aggrid.grid_options_builder import GridOptionsBuilder
import dash_ag_grid as dag              
from dash import Dash, html, dcc, Input, Output, State, no_update
import dash_bootstrap_components as dbc
import  openpyxl
import pandas as pd
import requests
import joblib
from io import BytesIO
from github import Github
import io
import base64
import pickle
from github import Github, UnknownObjectException
from google.oauth2 import service_account
import pyparsing
import gspread
#for cluster & mtbf
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.cluster import KMeans, MiniBatchKMeans
from sklearn.metrics import silhouette_score
import re
import os
from io import BytesIO
import numpy as np
from scipy.integrate import quad
from scipy.stats import norm
import matplotlib.pyplot as plt
import matplotlib.gridspec as gridspec
from docx import Document
from docx.shared import Inches
import warnings
#for chat
from langchain.embeddings.openai import OpenAIEmbeddings
from langchain.vectorstores import Chroma
from langchain.text_splitter import CharacterTextSplitter
from langchain.llms import OpenAI
from langchain.chains import RetrievalQA
from langchain.docstore.document import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.chains import RetrievalQAWithSourcesChain
from langchain.chains.question_answering import load_qa_chain
from langchain.prompts import PromptTemplate

warnings.filterwarnings("ignore")

# Create a connection object.
credentials = service_account.Credentials.from_service_account_info(
    st.secrets["gcp_service_account"],
    scopes=[
        "https://www.googleapis.com/auth/spreadsheets","https://www.googleapis.com/auth/drive"
    ],
)
client=gspread.authorize(credentials)

def run_query(query):
    rows = conn.execute(query, headers=1)
    rows = rows.fetchall()
    return rows

def validate_numeric(user_input):
    try:
        float(user_input)
        return True
    except ValueError:
        return False

def xlookup(lookup_value, lookup_array, return_array, if_not_found:str = ''):
    match_value = return_array.loc[lookup_array == lookup_value]
    if match_value.empty:
        return f'"{lookup_value}" not found!' if if_not_found == '' else if_not_found
    else:
        return match_value.tolist()[0]

def filter_dataframe(df: pd.DataFrame) -> pd.DataFrame:
    #get types_of_train sheet id
    types_of_train='1SX8ZLJWHZv_xHQ039gL02YsQt6szFhCMZylv9PshTao'
    #convert google sheet to csv for easy handling -> types_of_train (tot)
    tot_url=(f"https://docs.google.com/spreadsheets/d/{types_of_train}/export?format=csv")
    #create dataframe from csv
    tot_df=pd.read_csv(tot_url)
    #get unique Type of Train
    unq_tot=tot_df['Type of Train'].unique()
    # Define select_tot as the selected value from selectbox
    select_tot = st.selectbox('Pilih Jenis Kereta',unq_tot)
    #filter df based on type of train selection single selectbox
    filtered_tot = tot_df[tot_df['Type of Train'] == select_tot]
    #get unique level1 based on filtered tot df
    unq_filtered_tot=filtered_tot['MPG name'].unique()
    df = df.copy()
    to_filter_columns = st.multiselect("Main Product group apa sajakah yang akan Anda gunakan", unq_filtered_tot,unq_filtered_tot,key=1)
    modification_container = st.container()
    with modification_container:
        level2 = []
        level3 = [] 
        for i in to_filter_columns: 
            left, right = st.columns((1, 20))
            left.write("↳")
            with st.expander(f"Choose Sub Sub Product group (S-SPG) on {i}"):
                user_lvl2_input = right.multiselect(
                    f"Choose Sub Product group (SPG) on {i}",
                    df.loc[df['MPG name'] == i, 'SPG name'].unique(),
                    default=df.loc[df['MPG name'] == i, 'SPG name'].unique(),key=str(i)
                )
                level2 += user_lvl2_input
                for j in user_lvl2_input: 
                    left, right = st.columns((3, 20))
                    left.write("↳↳")
                    user_lvl3_input = right.multiselect(
                        f"Choose Sub Sub Product group (S-SPG) on {j}",
                        df.loc[df['SPG name'] == j, 'sub subproduct groups'].unique(),
                        default=df.loc[df['SPG name'] == j, 'sub subproduct groups'].unique(),key=str(i+j)
                    )
                    level3 += user_lvl3_input
    df = df[df['MPG name'].isin(to_filter_columns)]
    df = df[df['SPG name'].isin(level2)]
    df = df[df['sub subproduct groups'].isin(level3)]
    st.markdown("## Berikut komponen yang Anda pilih:")
    return df

def system_requirement():
    st.empty()
    st.title("Product Breakdown Structure")
    st.write(
        """This app streamlines the initial engingeeing process for a railway vehicle manufacturer by allowing users to select components based on the [BS EN 15380-2-2006 standard](https://drive.google.com/file/d/1O20tY4gVVmZVUSgSxAiYVSOxFs3tg48k/view?usp=share_link). It simplifies the initial steps of RAMS, especially on system selection, but please note that the app is not meant to fully cover the whole process. Human supervision is still necessary to ensure accuracy.
        """
    )
    tab1,tab2 = st.tabs(["Product Breakdown Picker 🍒","Product Breakdown Checker ✔️"])
    with tab1:
        #get sheet id
        sheet_id='1ikMHr99Z-IGOFcwK6a5On7soFtkGSG051LrKZvM-1sA'
        #convert google sheet to csv for easy handling
        csv_url=(f"https://docs.google.com/spreadsheets/d/{sheet_id}/export?format=csv")
        #create dataframe from csv
        df=pd.read_csv(csv_url)
        df=filter_dataframe(df)
        df.insert(2, "Qty/TS for level 1", "")
        df.insert(5, "Qty/TS for level 2", "")
        df.insert(7, "Qty/TS for level 3", "")
        df.insert(9, "Remark", "")
        gd=GridOptionsBuilder.from_dataframe(df)
        gd.configure_pagination(enabled=True)
        gd.configure_default_column(editable=True,groupable=True)
        gridoptions=gd.build()
        AgGrid(df,gridOptions=gridoptions, height=500, theme='alpine')
    with tab2:
        #get types_of_train sheet id
        types_of_train='1SX8ZLJWHZv_xHQ039gL02YsQt6szFhCMZylv9PshTao'
        #convert google sheet to csv for easy handling -> types_of_train (tot)
        tot_url=(f"https://docs.google.com/spreadsheets/d/{types_of_train}/export?format=csv")
        #create dataframe from csv
        tot_df=pd.read_csv(tot_url)
        #get unique Type of Train
        unq_tot=tot_df['Type of Train'].unique()
        # Define select_tot as the selected value from selectbox
        select_tot = st.selectbox('Pilih Jenis Kereta',unq_tot,key=str("select_tot"))
        #filter df based on type of train selection single selectbox
        filtered_tot = tot_df[tot_df['Type of Train'] == select_tot]
        #get unique level1 based on filtered tot df
        unq_filtered_tot=pd.DataFrame(filtered_tot['MPG name'].unique(),columns=['MPG name'])
        #get sheet id
        sheet_id='1ikMHr99Z-IGOFcwK6a5On7soFtkGSG051LrKZvM-1sA'
        #convert google sheet to csv for easy handling
        csv_url=(f"https://docs.google.com/spreadsheets/d/{sheet_id}/export?format=csv")
        #create dataframe from csv
        df=pd.read_csv(csv_url)       
        # left join with unq_filtered_tot
        joined_df = df.merge(unq_filtered_tot, on='MPG name', how='right')
        # Display the joined DataFrame
        gd=GridOptionsBuilder.from_dataframe(joined_df)
        gd.configure_default_column(editable=False,groupable=True)
        gridoptions=gd.build()
        AgGrid(joined_df,gridOptions=gridoptions, height=500, theme='alpine')
        uploaded_file2 = st.file_uploader("Upload a CSV or Excel file for the second dataset:", type=["csv", "xlsx"])
        if uploaded_file2 is not None:
            data2 = pd.read_csv(uploaded_file2) if uploaded_file2.type=='csv' else pd.read_excel(uploaded_file2)
        else:
            st.warning("Please upload a CSV or Excel file for the second dataset")
        if st.button("Compare Datasets"):
            # Get the data from the reference
            data1 = joined_df.iloc[:,:5]
            if data1 is not None and data2 is not None:
                # Filter the columns in dataset 2 that are not in dataset 1
                data2 = data2[data2.columns.intersection(data1.columns)]
                # Compare the datasets
                common = data1.merge(data2, on=data1.columns.tolist())
                not_in_data1 = data2[~data2.index.isin(common.index)]
                not_in_data2 = data1[~data1.index.isin(common.index)]
                # Create a DataFrame to display the results
                df = pd.DataFrame(columns=["Dataset 1", "Shared", "Not in Dataset 1", "Not in Dataset 2","Dataset 2"])
                for index in not_in_data1.index:
                    df = df.append({"Dataset 1": "", "Shared": "", "Not in Dataset 1": "✔", "Not in Dataset 2":"", "Dataset 2": not_in_data1.loc[index].values.tolist()}, ignore_index=True)
                for index in not_in_data2.index:
                    df = df.append({"Dataset 1": not_in_data2.loc[index].values.tolist(), "Shared": "", "Not in Dataset 1": "", "Not in Dataset 2": "✔", "Dataset 2":""},ignore_index=True)
                for index in common.index:
                    df = df.append({"Dataset 1": common.loc[index].values.tolist(), "Shared": "✔", "Not in Dataset 1": "", "Not in Dataset 2": "", "Dataset 2": common.loc[index].values.tolist()},ignore_index=True)                             
            # Display the DataFrame
            gd=GridOptionsBuilder.from_dataframe(df)
            gd.configure_default_column(editable=False,groupable=True)
            gridoptions=gd.build()
            AgGrid(df,gridOptions=gridoptions, height=500, theme='alpine')

def Supplier():
    st.empty()
    #get supplier id
    Suppplier='1_Gtz3x6yNI1qAvwaFdGSggfzrdf_6TtVvcL5Ny_6wOE'
    #convert google sheet to csv for easy handling -> types_of_train (tot)
    Suppplier_url=(f"https://docs.google.com/spreadsheets/d/{Suppplier}/export?format=csv")
    #create dataframe from csv
    Suppplier_df=pd.read_csv(Suppplier_url)
    unq_cat=Suppplier_df['Category'].unique()
    cat=st.selectbox("Pilih Category yang ada",unq_cat)
    unq_subcat = Suppplier_df[Suppplier_df['Category'].str.contains(cat)]['SubName Category'].unique()
    subcat=st.selectbox("Pilih Sub Category yang ada",unq_subcat)
    unq_subsubcat = Suppplier_df[Suppplier_df['Category'].str.contains(cat)&Suppplier_df['SubName Category'].str.contains(subcat)]['Sub SubName category'].unique()
    subsubcat=st.selectbox("Pilih Sub SubCategory yang ada",unq_subsubcat)
    #filter
    filtered_Suppplier = Suppplier_df[Suppplier_df['Category'].str.contains(cat)&Suppplier_df['SubName Category'].str.contains(subcat)&Suppplier_df['Sub SubName category'].str.contains(subsubcat)]
    st.write(f"{filtered_Suppplier.shape[0]} number of supplier found using keyword : {cat} & {subcat} & {subsubcat}")
    # Display the DataFrame
    gd=GridOptionsBuilder.from_dataframe(filtered_Suppplier )
    gd.configure_pagination(enabled=True)
    gd.configure_default_column(editable=False,groupable=True)
    gridoptions=gd.build()
    AgGrid(filtered_Suppplier ,gridOptions=gridoptions, height=500, theme='alpine')

def Standards():
    st.empty()
    # Define the URL of the file on the public GitHub repository
    file_url = 'https://raw.githubusercontent.com/bedy-kharisma/engineering/main/standards.pkl'
    # Download the file contents from the URL
    response = requests.get(file_url)
    if response.status_code == 200:
        content = BytesIO(response.content)
        standards = pd.read_pickle(content)
        keyword = st.text_input('Pilih keyword yang ingin Anda cari')
        #filter
        filtered_std = standards[standards['text'].str.contains(keyword, case=False)]
        standards_df=filtered_std[["location","name","id"]]
        if keyword!="":
            st.write(f"{standards_df.shape[0]} number of standards found using keyword : {keyword}")
        standards_df['link'] = standards_df['id'].apply(lambda x: f'<a target="_blank" href="https://drive.google.com/file/d/{x}/view">{x}</a>')
        standards_df = standards_df.to_html(escape=False)
        st.write(standards_df, unsafe_allow_html=True)
        
def FMECA():
    st.empty()
    uploaded_file2 = st.file_uploader("Upload a CSV or Excel file of Product Breakdown Structure:", type=["csv", "xlsx"])
    if uploaded_file2 is not None:
        data1 = pd.read_csv(uploaded_file2) if uploaded_file2.type=='csv' else pd.read_excel(uploaded_file2)
        #get types_of_train sheet id
        FMECA='18wTCYGDmtYUVJ5WBb_lrG-J6jY-O44v_Pl2G4mnyuxY'
        #convert google sheet to csv for easy handling -> types_of_train (tot)
        FMECA_url=(f"https://docs.google.com/spreadsheets/d/{FMECA}/export?format=csv")
        #create dataframe from csv
        FMECA_df=pd.read_csv(FMECA_url)
        # Filter the columns in dataset A to match the first 5 columns of dataset B
        columns_to_keep = FMECA_df.columns[:5]
        data1 = data1[data1.columns.intersection(FMECA_df.columns)]
        filter1 = FMECA_df["MPG name"].isin(data1["MPG name"])
        filter2 = FMECA_df["SPG name"].isin(data1["SPG name"])
        filter3 = FMECA_df["sub subproduct groups"].isin(data1["sub subproduct groups"])
        merged_df=FMECA_df[filter1&filter2&filter3]
    else:
        st.warning("Please upload a CSV or Excel file for the second dataset") 
    if st.button("Generate initial FMECA") and uploaded_file2 is not None:
        # Display the DataFrame
        gd=GridOptionsBuilder.from_dataframe(merged_df)
        gd.configure_pagination(enabled=True)
        gd.configure_default_column(editable=False,groupable=True)
        gridoptions=gd.build()
        AgGrid(merged_df,gridOptions=gridoptions, height=500, theme='alpine')

def failure():
    st.empty()
    #get sheet id
    sheet_id='1yYY6kEVkBRRdmNcGG1cG2F3gbPA_OZJ5rUTWahx5X-U'
    #convert google sheet to csv for easy handling
    csv_url=(f"https://docs.google.com/spreadsheets/d/{sheet_id}/export?format=csv")
    #create dataframe from csv
    df=pd.read_csv(csv_url)
    unq_mainsys=df['Main System'].unique()
    df['Jumlah kegagalan']= df['Jumlah kegagalan'].astype(int)
    st.markdown("# Failure Rate Calculator")
    st.write("This failure rate calcuator is populated using database saved in this [google sheet](https://docs.google.com/spreadsheets/d/1yYY6kEVkBRRdmNcGG1cG2F3gbPA_OZJ5rUTWahx5X-U/edit?usp=sharing)")
    mainsys = st.selectbox('Pilih Main System apa yang akan Anda hitung',unq_mainsys)
    #filter df based on main system selection single selectbox
    filtered_df = df[df['Main System'] == mainsys]
    #get unique project name based on filtered df
    unq_proj=filtered_df['Project Name'].unique()
    proj = st.multiselect('Data dari Project saja yang akan Anda gunakan',unq_proj,unq_proj)
    #filter df based on multiple selection project name
    filtered_df = filtered_df[filtered_df['Project Name'].isin(proj)]
    #get unique vendor based on above filtered df
    unq_vendor=filtered_df['Vendor Name'].unique()
    vendor=st.multiselect('Data dari Vendor saja yang akan Anda gunakan',unq_vendor,unq_vendor)
    #filter based on multiple select vendor name
    filtered_df = filtered_df[filtered_df['Vendor Name'].isin(vendor)]
    # Create a new dataframe with the duplicates removed
    df_no_duplicates = filtered_df.drop_duplicates(subset='Project Name', inplace=False)
    # sum the values in the 'operating hours per year' column
    total_ophours = df_no_duplicates['Operating Hours per Year'].sum()
    # Assign the sum back to the original dataframe
    filtered_df['Total Operating Hours'] = total_ophours.astype(int)
    #calculate total quantity
    filtered_df['Total Quantity'] = filtered_df.groupby('Item Name')['Quantity all Trainset'].transform('sum')
    filtered_df['Total Quantity']= filtered_df['Total Quantity'].astype(int)
    #calculate t = Total komponen*OH
    filtered_df['Total komponen*OH']= filtered_df['Total Quantity']*filtered_df['Total Operating Hours']
    filtered_df['Total komponen*OH']= filtered_df['Total komponen*OH']
    # Define the function
    def failure_rate(row):
        if row['Total Quantity'] == 0 or row['Total Operating Hours'] == 0:
            return 0
        if row['Jumlah kegagalan'] > 0:
            return row['Jumlah kegagalan'] / row['Total komponen*OH']
        else:
            return 1 / (row['Total Quantity'] * row['Total Operating Hours'])
    # Apply the function to the dataframe
    filtered_df['Failure Rate'] = filtered_df.apply(failure_rate, axis=1)
    #change into scientific format
    def to_scientific(x):
        return '{:.2e}'.format(x)
    filtered_df['Failure Rate'] = filtered_df['Failure Rate'].apply(to_scientific)
    #show only unique item
    unique_df = filtered_df.drop_duplicates(subset=['Item Ref', 'Item Name'])
    unique_df  = unique_df .reset_index().drop(columns='index')
    st.dataframe(unique_df[['Item Ref', 'Item Name','Total Operating Hours','Total Quantity','Jumlah kegagalan','Failure Rate']])

def FBS():
    st.empty()
    st.title("Function Breakdown Structure")
    st.write(
        """This app streamlines the initial engingeeing process for a railway vehicle manufacturer by allowing users to select functions based on the [BS EN 15380-4-2013 standard](https://drive.google.com/file/d/19Wmq1jLGlQdNZL9UpnUgL80Ec5c-gC1M/view?usp=share_link). It simplifies the initial steps of RAMS, such as system requirements, selection, and design, but please note that the app is not meant to fully cover the whole process. Human supervision is still necessary to ensure accuracy.
        """
    )
    tab1,tab2 = st.tabs(["Function Breakdown Picker 🍒","Function Breakdown Checker ✔️"])
    with tab1:
        #get sheet id level1-3 FBS
        sheet_id='1yEtUmBuxRewIiclGothFn9IaZdgm458owf8V_ZJnYQk'
        #convert google sheet to csv for easy handling
        csv_url=(f"https://docs.google.com/spreadsheets/d/{sheet_id}/export?format=csv")
        #create dataframe from csv
        fbs_df=pd.read_csv(csv_url,on_bad_lines='skip')
        ##################
        to_filter_columns = st.multiselect("Function Level 1 apa sajakah yang akan Anda gunakan", fbs_df['level 1 Function'].unique(),fbs_df['level 1 Function'].unique(),key=1)
        modification_container = st.container()
        with modification_container:
            level2 = []
            level3 = []
            for i in to_filter_columns: 
                if pd.isna(i):
                    continue
                left, right = st.columns((1, 20))
                left.write("↳")
                with st.expander(f"Choose Level 3 Function group on {i}"):
                    user_lvl2_input = right.multiselect(
                        f"Choose Level 2 Function group on {i}",
                        fbs_df.loc[fbs_df['level 1 Function'] == i, 'level 2 Function'].dropna().unique(),
                        default=fbs_df.loc[fbs_df['level 1 Function'] == i, 'level 2 Function'].dropna().unique(),key=str(i)
                    )
                    level2 += user_lvl2_input
                    for j in user_lvl2_input: 
                        if pd.isna(j):
                            continue
                        left, right = st.columns((3, 20))
                        left.write("↳↳")
                        user_lvl3_input = right.multiselect(
                            f"Choose Level 3 Function on {j}",
                            fbs_df.loc[fbs_df['level 2 Function'] == j, 'level 3 Function'].dropna().unique(),
                            default=fbs_df.loc[fbs_df['level 2 Function'] == j, 'level 3 Function'].dropna().unique(),key=str(i) + str(j)
                        )
                        level3 += user_lvl3_input
        fbs_df = fbs_df[fbs_df['level 1 Function'].isin(to_filter_columns)]
        fbs_df = fbs_df[fbs_df['level 2 Function'].isin(level2)]
        fbs_df = fbs_df[fbs_df['level 3 Function'].isin(level3)]
        st.markdown("## Berikut Functions yang Anda pilih:")
        gd=GridOptionsBuilder.from_dataframe(fbs_df)
        gd.configure_default_column(editable=True,groupable=True)
        gridoptions=gd.build()
        AgGrid(fbs_df,gridOptions=gridoptions, height=800, theme='alpine')
    with tab2:
        #get sheet id level1-3 FBS
        sheet_id='1yEtUmBuxRewIiclGothFn9IaZdgm458owf8V_ZJnYQk'
        #convert google sheet to csv for easy handling
        csv_url=(f"https://docs.google.com/spreadsheets/d/{sheet_id}/export?format=csv")
        #create dataframe from csv
        fbs_df=pd.read_csv(csv_url,on_bad_lines='skip')
        ##################
        # Display the joined DataFrame
        gd=GridOptionsBuilder.from_dataframe(fbs_df)
        gd.configure_default_column(editable=False,groupable=True)
        gridoptions=gd.build()
        AgGrid(fbs_df,gridOptions=gridoptions, height=500, theme='alpine')
        uploaded_file2 = st.file_uploader("Upload a CSV or Excel file for the second dataset:", type=["csv", "xlsx"])
        if uploaded_file2 is not None:
            data2 = pd.read_csv(uploaded_file2) if uploaded_file2.type=='csv' else pd.read_excel(uploaded_file2)
        else:
            st.warning("Please upload a CSV or Excel file for the second dataset")
        if st.button("Compare Datasets"):
            # Get the data from the reference
            data1 = fbs_df.iloc[:,:5]
            if data1 is not None and data2 is not None:
                # Filter the columns in dataset 2 that are not in dataset 1
                data2 = data2[data2.columns.intersection(data1.columns)]
                # Compare the datasets
                common = data1.merge(data2, on=data1.columns.tolist())
                not_in_data1 = data2[~data2.index.isin(common.index)]
                not_in_data2 = data1[~data1.index.isin(common.index)]
                # Create a DataFrame to display the results
                df = pd.DataFrame(columns=["Dataset 1", "Shared", "Not in Dataset 1", "Not in Dataset 2","Dataset 2"])
                for index in not_in_data1.index:
                    df = df.append({"Dataset 1": "", "Shared": "", "Not in Dataset 1": "✔", "Not in Dataset 2":"", "Dataset 2": not_in_data1.loc[index].values.tolist()}, ignore_index=True)
                for index in not_in_data2.index:
                    df = df.append({"Dataset 1": not_in_data2.loc[index].values.tolist(), "Shared": "", "Not in Dataset 1": "", "Not in Dataset 2": "✔", "Dataset 2":""},ignore_index=True)
                for index in common.index:
                    df = df.append({"Dataset 1": common.loc[index].values.tolist(), "Shared": "✔", "Not in Dataset 1": "", "Not in Dataset 2": "", "Dataset 2": common.loc[index].values.tolist()},ignore_index=True)                                   
            # Display the DataFrame
            gd=GridOptionsBuilder.from_dataframe(df)
            gd.configure_default_column(editable=False,groupable=True)
            gridoptions=gd.build()
            AgGrid(df,gridOptions=gridoptions, height=500, theme='alpine')

def Matcod():
    sheet_id='1uFfcegQlGi6vKtyuhq_RxDKRJ26fw_bGod2Lic5Bjy8'
    #convert google sheet to csv for easy handling
    csv_url=(f"https://docs.google.com/spreadsheets/d/{sheet_id}/export?format=csv")
    #create dataframe from csv
    database_df=pd.read_csv(csv_url,on_bad_lines='skip')   
    pickle_file = 'TB1_df.pkl'
    file_url = 'https://raw.githubusercontent.com/bedy-kharisma/engineering/main/'+ pickle_file
    response = requests.get(file_url)
    if response.status_code == 200:
        content = BytesIO(response.content) 
        TB1_df=pd.read_pickle(content) 
    pickle_file = 'TB2_df.pkl'   
    file_url = 'https://raw.githubusercontent.com/bedy-kharisma/engineering/main/'+ pickle_file
    response = requests.get(file_url)
    if response.status_code == 200:
        content = BytesIO(response.content) 
        TB2_df=pd.read_pickle(content) 
    pickle_file = 'Fastening_df.pkl' 
    file_url = 'https://raw.githubusercontent.com/bedy-kharisma/engineering/main/'+ pickle_file
    response = requests.get(file_url)
    if response.status_code == 200:
        content = BytesIO(response.content) 
        Fastening_df=pd.read_pickle(content) 
    pickle_file = 'maincom_df.pkl'
    file_url = 'https://raw.githubusercontent.com/bedy-kharisma/engineering/main/'+ pickle_file
    response = requests.get(file_url)
    if response.status_code == 200:
        content = BytesIO(response.content) 
        maincom_df=pd.read_pickle(content) 
    pickle_file = 'sw_df.pkl' 
    file_url = 'https://raw.githubusercontent.com/bedy-kharisma/engineering/main/'+ pickle_file
    response = requests.get(file_url)
    if response.status_code == 200:
        content = BytesIO(response.content)
        sw_df=pd.read_pickle(content) 
    pickle_file = 'el_df.pkl'
    file_url = 'https://raw.githubusercontent.com/bedy-kharisma/engineering/main/'+ pickle_file
    response = requests.get(file_url)
    if response.status_code == 200:
        content = BytesIO(response.content)
        el_df=pd.read_pickle(content) 
    pickle_file = 'brake_df.pkl'
    file_url = 'https://raw.githubusercontent.com/bedy-kharisma/engineering/main/'+ pickle_file
    response = requests.get(file_url)
    if response.status_code == 200:
        content = BytesIO(response.content)
        brake_df=pd.read_pickle(content) 
    pickle_file = 'bogie_df.pkl'
    file_url = 'https://raw.githubusercontent.com/bedy-kharisma/engineering/main/'+ pickle_file
    response = requests.get(file_url)
    if response.status_code == 200:
        content = BytesIO(response.content)
        bogie_df=pd.read_pickle(content) 
    pickle_file = 'coupler_df.pkl'    
    file_url = 'https://raw.githubusercontent.com/bedy-kharisma/engineering/main/'+ pickle_file
    response = requests.get(file_url)
    if response.status_code == 200:
        content = BytesIO(response.content)
        coupler_df=pd.read_pickle(content) 
    pickle_file = 'interior_df.pkl'
    file_url = 'https://raw.githubusercontent.com/bedy-kharisma/engineering/main/'+ pickle_file
    response = requests.get(file_url)
    if response.status_code == 200:
        content = BytesIO(response.content)
        interior_df=pd.read_pickle(content) 
    pickle_file = 'piping_df.pkl'
    file_url = 'https://raw.githubusercontent.com/bedy-kharisma/engineering/main/'+ pickle_file
    response = requests.get(file_url)
    if response.status_code == 200:
        content = BytesIO(response.content)
        piping_df=pd.read_pickle(content)
    pickle_file = 'cons_df.pkl'
    file_url = 'https://raw.githubusercontent.com/bedy-kharisma/engineering/main/'+ pickle_file
    response = requests.get(file_url)
    if response.status_code == 200:
        content = BytesIO(response.content)
        cons_df=pd.read_pickle(content)
    pickle_file = 'tools_df.pkl'
    file_url = 'https://raw.githubusercontent.com/bedy-kharisma/engineering/main/'+ pickle_file
    response = requests.get(file_url)
    if response.status_code == 200:
        content = BytesIO(response.content)
        tools_df=pd.read_pickle(content)
    pickle_file = 'raw_df.pkl' 
    file_url = 'https://raw.githubusercontent.com/bedy-kharisma/engineering/main/'+ pickle_file
    response = requests.get(file_url)
    if response.status_code == 200:
        content = BytesIO(response.content)
        raw_df=pd.read_pickle(content)
    pickle_file = 'spare_df.pkl'
    file_url = 'https://raw.githubusercontent.com/bedy-kharisma/engineering/main/'+ pickle_file
    response = requests.get(file_url)
    if response.status_code == 200:
        content = BytesIO(response.content)
        spare_df=pd.read_pickle(content)
    pickle_file = 'facilities_df.pkl'
    file_url = 'https://raw.githubusercontent.com/bedy-kharisma/engineering/main/'+ pickle_file
    response = requests.get(file_url)
    if response.status_code == 200:
        content = BytesIO(response.content)
        facilities_df=pd.read_pickle(content)
    st.title("MATERIAL CODE")
    tab1, tab2 = st.tabs(["Request", "Verification"])
    with tab1:
        # Display the DataFrame
        gd=GridOptionsBuilder.from_dataframe(database_df )
        gd.configure_pagination(enabled=True)
        gd.configure_default_column(editable=False,groupable=True)
        gridoptions=gd.build()
        AgGrid(database_df ,gridOptions=gridoptions, height=500, theme='alpine')
        st.write("### If you are sure that what you are looking for is not listed there, please fill up the entry form below:")
        #get unique
        unq_TB1=TB1_df['NAMA'].unique()
        # Define select as the selected value from selectbox
        select_TB1 = st.selectbox('Choose TB1 Code',unq_TB1)
        select_TB1 = xlookup(select_TB1, TB1_df['NAMA'],TB1_df['CODE'])
        st.write(select_TB1)
        #filter TB2 based on type of train selection single selectbox
        filtered_TB2 = TB2_df[TB2_df['CODE_TB1'].str.contains(select_TB1, na=False)]
        unq_TB2=filtered_TB2['NAMA'].unique()
        # Define select as the selected value from selectbox
        select_TB2 = st.selectbox('Choose TB2 Code',unq_TB2)
        select_TB2 = xlookup(select_TB2, TB2_df['NAMA'],TB2_df['CODE_TB2'])
        code=select_TB1+select_TB2
        st.write(code)
        df_dict = {"A54": el_df,"A52": el_df,"B39": Fastening_df, "B40": maincom_df, "B54": sw_df, "B52": el_df, "B47": brake_df, "B48": bogie_df, "B49": coupler_df, "B50": interior_df, "B51": piping_df, "D29": cons_df, "D30": cons_df, "D31": cons_df, "D32": cons_df, "D33": cons_df, "D61": cons_df, "D62": cons_df, "D63": cons_df, "D64": cons_df, "D65": cons_df, "D66": cons_df, "D67": cons_df, "D68": cons_df, "D69": cons_df, "D70": cons_df, "D71": cons_df, "D72": cons_df, "D73": cons_df, "D74": cons_df, "D75": cons_df, "D76": cons_df, "D80": cons_df, "D82": cons_df, "D83": cons_df, "D84": cons_df, "D98": cons_df, "D99": cons_df, "C71": tools_df, "C77": tools_df, "C78": tools_df, "C79": tools_df, "C85": tools_df, "C86": tools_df, "C87": tools_df, "C88": tools_df, "C89": tools_df, "C90": tools_df, "C91": tools_df, "C92": tools_df, "C93": tools_df, "C94": tools_df, "C95": tools_df, "C96": tools_df, "A01": raw_df, "A04": raw_df, "A09": raw_df, "A10": raw_df, "A11": raw_df, "A12": raw_df, "A13": raw_df, "A14": raw_df, "A15": raw_df, "A16": raw_df, "A17": raw_df, "A18": raw_df, "A19": raw_df, "A20": raw_df, "A21": raw_df, "A22": raw_df, "A23": raw_df, "A24": raw_df, "A25": raw_df, "B98": spare_df, "D98": spare_df, "E97": facilities_df}
        skip_codes = ["B37", "B41", "B42", "B43", "B44","B45", "B46"]
        if code not in skip_codes:
            selected_df = df_dict[code]
            if selected_df is not None:
                gd=GridOptionsBuilder.from_dataframe(selected_df)
                gd.configure_pagination(enabled=False)
                gd.configure_default_column(editable=False,groupable=True)
                gridoptions=gd.build()
                AgGrid(selected_df ,gridOptions=gridoptions, height=500, theme='alpine', fit_columns_on_grid_load=True)
                unq_selected_df=selected_df['DESCRIPTION'].unique()
                select_TB3 = st.selectbox('Choose TB3 Description',unq_selected_df)
                select_TB3 = xlookup(select_TB3, selected_df['DESCRIPTION'],selected_df['CODE'])
                code=select_TB1+select_TB2+select_TB3
                st.write(code)
                df = database_df[database_df["Kode Material"].str[:7].isin([code])]
        else:
                code=select_TB1+select_TB2
                df = database_df[database_df["Kode Material"].str[:3].isin([code])]
        # Display the DataFrame
        gd=GridOptionsBuilder.from_dataframe(df)
        gd.configure_pagination(enabled=False)
        gd.configure_default_column(editable=False,groupable=True)
        gridoptions=gd.build()
        AgGridaggrid = AgGrid(df, gridOptions=gridoptions, height=500, theme='alpine', 
                         data_return_mode=DataReturnMode.AS_INPUT, update_on='VALUE_CHANGED',        
                         enable_enterprise_modules=True, update_mode=GridUpdateMode.SELECTION_CHANGED,
                         allow_unsafe_jscode=True, fit_columns_on_grid_load=True)
        st.write("### If you are sure that what you are looking for is not listed there, please fill up the entry form below:")
        user_input = st.text_input("Insert unique number id", "",max_chars=7, key="input")
        if validate_numeric(user_input) and (len(code+user_input)<=12):
            st.write("New Material Code :")
            value_to_check=code+user_input
            if value_to_check in database_df["Kode Material"].unique():
                st.write(f'{value_to_check} is not unique in the material code database.')
            else:
                st.write(f'{value_to_check} is unique in the material code database.')
                with st.form("entry",clear_on_submit=True):
                    deskripsi = st.text_input("Insert description")
                    spec = st.text_input("Insert specification")
                    uom = st.text_input("Insert UoM")
                    requester = st.text_input("Insert Requester ID")
                    submit= st.form_submit_button("Submit")
                if submit:
                    database_df = database_df.append({'Kode Material': code+user_input, 'Deskripsi': deskripsi, 'Specification':   spec,'UoM':  uom,'Requester':   requester, 'Verification Status': "Unverified"}, ignore_index=True)    
                    database_df = database_df.astype(str)
                    sheet_url = st.secrets["private_gsheets_url"]
                    sheet=client.open("database").sheet1
                    sheet.update([database_df.columns.values.tolist()]+database_df.values.tolist())
                    st.success('New Material Code has been generated, Contact your EIM to verify it')
        else:
            st.write('Please enter a numeric value only & make sure the length is <= 12 characters')
    with tab2:
        password=st.text_input("Insert admin password","",type="password")
        if password == "admin":
            funct=st.radio(label="Functions:", options=['Edit','Delete'])
            st.write("or edit directly on this google sheet [google sheet](https://docs.google.com/spreadsheets/d/1uFfcegQlGi6vKtyuhq_RxDKRJ26fw_bGod2Lic5Bjy8/edit?usp=sharing)")
            if funct =='Delete':
                js=JsCode("""
                    function(e) {
                        let api = e.api;
                        let sel = api.getSelectedRows();
                        api.applyTransaction({remove: sel})
                    };
                """)
                sheet_id='1uFfcegQlGi6vKtyuhq_RxDKRJ26fw_bGod2Lic5Bjy8'
                #convert google sheet to csv for easy handling
                csv_url=(f"https://docs.google.com/spreadsheets/d/{sheet_id}/export?format=csv")
                #create dataframe from csv
                database_df=pd.read_csv(csv_url,on_bad_lines='skip')
                # Display the DataFrame
                gd=GridOptionsBuilder.from_dataframe(database_df)
                gd.configure_selection(selection_mode='single',use_checkbox=True)
                gd.configure_grid_options(onRowSelected=js,pre_selected_rows=[])
                gd.configure_pagination(enabled=False)
                gd.configure_default_column(editable=True,groupable=True)
                gridoptions=gd.build()
                aggrid = AgGrid(database_df, gridOptions=gridoptions, height=500, theme='alpine', 
                         data_return_mode=DataReturnMode.AS_INPUT, update_on='VALUE_CHANGED',        
                         enable_enterprise_modules=True, update_mode=GridUpdateMode.SELECTION_CHANGED,
                         allow_unsafe_jscode=True, fit_columns_on_grid_load=True)
                # Update the original DataFrame
                data=aggrid['data']
                database_df=pd.DataFrame(data)
                database_df = database_df.astype(str)
                sheet_url = st.secrets["private_gsheets_url"]
                sheet=client.open("database").sheet1
                sheet.update([database_df.columns.values.tolist()]+database_df.values.tolist())
                st.info("Total rows :"+str(len(database_df)))       
            if funct =='Edit':
                # Display the DataFrame
                gd=GridOptionsBuilder.from_dataframe(database_df)
                gd.configure_pagination(enabled=False)
                gd.configure_default_column(editable=True,groupable=True)
                gridoptions=gd.build()
                aggrid = AgGrid(database_df ,gridOptions=gridoptions, height=500, theme='alpine', 
                                data_return_mode=DataReturnMode.AS_INPUT, update_on='VALUE_CHANGED',
                                fit_columns_on_grid_load=True)
                # Update the original DataFrame
                data=aggrid['data']
                database_df=pd.DataFrame(data)
                database_df = database_df.astype(str)
                sheet_url = st.secrets["private_gsheets_url"]
                sheet=client.open("database").sheet1
                sheet.update([database_df.columns.values.tolist()]+database_df.values.tolist())
                st.info("Total rows :"+str(len(database_df)))

# FITTINGS
def ExponentialFitting(Data):
    Data = np.sort(Data)
    r = len(Data)
    biased_lambda = 1 / np.mean(Data)
    lam = (r - 1) / r * biased_lambda
    return lam

def gumbelfitting(Data, n):
    sorted_data = np.sort(Data)
    mu = np.mean(sorted_data[-n:])
    sigma = np.std(sorted_data[-n:])
    return mu, sigma

def LogNormalFitting(Data):
    log_data = np.log(Data)
    mu = np.mean(log_data)
    stdev = np.std(log_data)
    shape = mu
    scale = stdev
    return shape, scale

def NormalDistributionFitting(Data):
    mu = np.mean(Data)
    sigma = np.std(Data)
    return mu, sigma

def weibull_fitting_lse(Data):
    Data = np.sort(Data)
    r = len(Data)  
    w = np.zeros(r)
    y = np.zeros(r)  
    for i in range(r):
        Fx = (i + 1) / (r + 1)
        w[i] = ((1 - Fx) * np.log(1 - Fx))**2
        y[i] = np.log(-np.log(1 - Fx))   
    wlog = w * np.log(Data)
    wlogy = wlog * y
    wy = w * y
    wlog2 = wlog * np.log(Data) 
    Num = np.sum(w) * np.sum(wlogy) - np.sum(wlog) * np.sum(wy)
    Den = np.sum(w) * np.sum(wlog2) - np.sum(wlog)**2
    shape = Num / Den
    Numerator = np.sum(wy) - shape * np.sum(wlog)
    Denumerator = shape * np.sum(w)
    scale = np.exp(-Numerator / Denumerator)
    return shape, scale

def WeibullFittingMLE(Data):
    Data = np.sort(Data)
    r = len(Data)
    third = np.sum(np.log(Data)) / r
    shape = None
    for i in np.arange(0.01, 10, 0.001):
        first = 1 / i
        second = np.sum((Data**i) * np.log(Data)) / np.sum(Data**i)
        if first - second + third < 0.0001:
            shape = i
            break
    scale = (np.sum(Data**shape) / r) ** (1 / shape)  
    return shape, scale
#----END FITTING

# Distribution tests
def normal_distribution_test(data):
    W = 0
    n = len(data)
    nh = int(np.floor(n / 2))
    # Sorting the data
    Sorted_Data = np.sort(data)
    # Mean Calculation
    Data_Avg = np.mean(data)
    # S squared calculation
    S_Squared = 0
    for i in range(len(data)):
        S_Squared = S_Squared + (Sorted_Data[i] - Data_Avg)**2
    # a coefficient calculation
    if n == 3:
        a = np.sqrt(0.5)  # exact
    else:
        # get an initial estimate of a
        from scipy.stats import norm
        m = -norm.ppf(((np.arange(1, nh+1) - 0.375) / (n + 0.25)))
        msq = 2 * np.sum(m**2)
        mrms = np.sqrt(msq)
        # correction factors for initial elements of a (derived to be good approximations for 4 <= n <= 1000)
        rsn = 1 / np.sqrt(n)
        if n <= 5:
            polyval_coeffs = [-2.706056, 4.434685, -2.07119, -0.147981, 0.221157, 0]
            ac1 = m / mrms + np.polyval(polyval_coeffs, rsn)
            phi = (msq - 2 * m**2) / (1 - 2 * ac1**2)
            a = np.concatenate([ac1.reshape(-1, 1), m[1:] / np.sqrt(phi).reshape(-1, 1)])
        else:
            polyval_coeffs = [-3.582633, 5.682633, -1.752461, -0.293762, 0.042981, 0]
            ac1 = m[0] / mrms + np.polyval(polyval_coeffs, rsn)
            ac2 = m[1] / mrms + np.polyval(polyval_coeffs, rsn)
            phi = (msq - 2 * m[0]**2 - 2 * m[1]**2) / (1 - 2 * ac1**2 - 2 * ac2**2)
            a = np.concatenate(([ac1, ac2], m[2:] / np.sqrt(phi)))
    W = np.sum(a * (Sorted_Data[n - np.arange(nh) - 1] - Sorted_Data[:nh])) ** 2 / float(S_Squared)
    # p value for the W statistic being as small as it is for a normal distribution
    if n == 3:
        pval = (6 / np.pi) * (np.arcsin(np.sqrt(W)) - np.arcsin(np.sqrt(0.75)))
    elif n <= 11:
        gamma = 0.459 * n - 2.273
        w = -np.log(gamma - np.log(1 - W))
        mu = np.polyval([-0.0006714, 0.025054, -0.39978, 0.5440], gamma)
        sigma = np.exp(np.polyval([-0.0020322, 0.062767, -0.77857, 1.3822], n))
        pval = norm.cdf((mu - w) / sigma)
    else:
        nl = np.log(n)
        w = np.log(1 - W)
        mu = np.polyval([0.0038915, -0.083751, -0.31082, -1.5861], nl)
        sigma = np.exp(np.polyval([0.0030302, -0.082676, -0.4803], nl))
        pval = norm.cdf((mu - w) / sigma)
    return W, pval

def exponential_test(data, r):
    if r is None and len(data) > 0:
        r = len(data)
    # denominator calculation
    den = 1 + (r + 1) / (6 * r)
    # numerator calculation
    num = 2 * r * (np.log((1 / r) * np.sum(data)) - (np.sum(np.log(data)) / r))
    B = num / den
    return B

def weibull_test(data, n):
    Sorted_Data = np.sort(data)
    r = len(data)
    Z = np.zeros(len(data))
    M = np.zeros(len(data)-1)
    # Calculation of Z
    for i in range(1, r+1):
        Z[i-1] = np.log(-np.log(1-((i-0.5)/(n+2.5))))
    M = np.diff(Z)    
    # Calculation of k
    k1 = int(np.floor(r/2))
    k2 = int(np.floor((r-1)/2))
    # Calculation of numerator and denominator for M
    num = 0
    den = 0
    for i in range(k1+1, r-1):
        num = num + (np.log(Sorted_Data[i+1]) - np.log(Sorted_Data[i]))/M[i]
    num = k1 * num
    for i in range(k1):
        den = den + (np.log(Sorted_Data[i+1]) - np.log(Sorted_Data[i]))/M[i]
    den = k2 * den
    M = num / den
    dof1 = 2 * k1
    dof2 = 2 * k2
    return M, dof1, dof2
#--- END Distribution tests

# Cases function
def cases(data, distribution, x):
    RMTTF = 0
    R = np.zeros(len(x))
    hr = np.zeros(len(x))
    PDF = np.zeros(len(x))
    if distribution == "Normal":
        mu, sigma = NormalDistributionFitting(data)
        MU = mu
        SIG = sigma
        f = lambda t: 1/(SIG*np.sqrt(2*np.pi))*np.exp(-(1/2)*((t-MU)/SIG)**2)
        MTTF = mu
        for i in range(len(x)):
            R[i],_= quad(f, x[i], np.inf)
            hr[i] = f(x[i])/R[i]
            PDF[i] = f(x[i])
            if abs(MTTF - x[i]) <= 0.99:
                RMTTF = R[i]  
    elif distribution == "Log-Normal":
        x[0] = 1
        location, scale = LogNormalFitting(data)
        LOC = round(location)
        SCA = round(scale)
        f = lambda t: (1/(t*SCA*np.sqrt(2*np.pi)))*np.exp(-(1/2)*((np.log(t)-LOC)/(SCA))**2)
        mttf = lambda t: (t/(t*LOC*np.sqrt(2*np.pi)))*np.exp(-(1/2)*((np.log(t)-LOC)/(SCA))**2)
        MTTF,_ = quad(mttf, 0, np.inf)
        for i in range(len(x)):
            R[i],_ = quad(f, x[i], np.inf)
            hr[i] = f(x[i])/R[i]
            PDF[i] = f(x[i])
            if abs(MTTF - x[i]) <= 0.99:
                RMTTF = R[i]  
    elif distribution == "Exponential":
        lambda_ = ExponentialFitting(data)
        f = lambda t: lambda_*np.exp(-lambda_*t)
        mttf = lambda t: t*lambda_*np.exp(-lambda_*t)
        MTTF,_ = quad(mttf, 0.01, np.inf)
        for i in range(len(x)):
            R[i],_ = quad(f, x[i], np.inf)
            hr[i] = f(x[i])/R[i]
            PDF[i] = f(x[i])
            if abs(MTTF - x[i]) <= 0.99:
                RMTTF = R[i]
        RMTTF = np.exp(-(lambda_*MTTF))   
    elif distribution == "Weibull":
        shape, scale = WeibullFittingMLE(data)
        f = lambda t: (shape/scale)*((t/scale)**(shape-1))*np.exp(-(t/scale)**(shape-1))
        Gamm = lambda t: t**((1/shape)+1)*np.exp(-t)
        MTTF,_ = quad(Gamm, 0, np.inf)
        MTTF=scale*MTTF
        for i in range(len(x)):
            R[i] = np.exp(-(x[i]/scale)**shape)
            hr[i] = f(x[i])/R[i]
            PDF[i] = f(x[i])
            if abs(MTTF - x[i]) <= 0.99:
                RMTTF = R[i]
        print("Shape:", shape)
        print("Scale:", scale)  
    else:
        print("Error: Invalid distribution input")
    return hr, R, PDF, MTTF, RMTTF
#--- END CASES

#FIGS
def plot_results(hr, R, PDF, MTTF, RMTTF,data,distribution):
    if MTTF <= 10000:
        digit=10000
        t = np.arange(10001)
    else:
        num_digit=len(str(round(MTTF,0)))-1
        digit=10**num_digit
        t = np.arange((10**num_digit)+1)
        hr, R, PDF, MTTF, RMTTF = cases(data, distribution, t)
    fig = plt.figure(figsize=(16, 9))
    gs = gridspec.GridSpec(3, 1, height_ratios=[1, 1, 1.2], hspace=0.5)
    # Plot 1: Probability Density Function
    ax1 = fig.add_subplot(gs[0])
    ax1.plot(t, PDF)
    ax1.grid(True)
    ax1.set_xlabel("Time (hours)")
    ax1.set_ylabel("PDF (-)")
    ax1.set_title("Probability Density Function")
    # Plot 2: Reliability
    ax2 = fig.add_subplot(gs[1])
    ax2.plot(t, R)
    ax2.plot(MTTF, RMTTF, "ok")
    ax2.set_xlabel("Time (hours)")
    ax2.set_ylabel("Reliability (-)")
    ax2.set_title("Reliability plot")
    ax2.grid(True)
    ax2.set_ylim([0, 1.01])
    xticks = [0, 0.25*digit, 0.5*digit, 0.75*digit, digit]
    xticklabels = ["0", str(0.25*digit), str(0.5*digit), str(0.75*digit), str(digit)]
    insert_pos = 0
    for i in range(len(xticks)-1):
        if MTTF >= xticks[i] and MTTF <= xticks[i+1]:
            insert_pos = i + 1
            break
    xticks.insert(insert_pos, MTTF)
    xticklabels.insert(insert_pos, str(round(MTTF,2)))
    ax2.set_xticks(xticks)
    ax2.set_xticklabels(xticklabels)
    ax2.legend(["Reliability", "Predicted Failure"])
    # Plot 3: Hazard Rate
    ax3 = fig.add_subplot(gs[2])
    ax3.plot(t, hr)
    ax3.set_xlabel("Time (hours)")
    ax3.set_ylabel("Hazard rate (-)")
    ax3.set_title("Hazard rate plot")
    ax3.grid(True)
    return fig
#--- END FIGS

def test(df, distribution, doc):
    distribution=distribution
    # Perform distribution tests and display the results    
    Raw = df
    row = len(Raw)
    column = len(Raw.columns)
    Data = Raw.iloc[:, column - 2:column].values
    final = np.sum(Data[:, 0])
    final=int(final)
    output = np.zeros(final)
    output_data = []
    for row in Data:
        count = row[0]  # Number of repetitions
        value = row[1]  # Value to be repeated
        for _ in range(count):
            output_data.append(value)
    output_array = np.array(output_data)
    data = output_array
    sorted_data = np.sort(data)
    log_data = np.log(sorted_data)
    st.write("The minimum value of p value for alpha = 0.05 is 0.05")
    st.write("If the tested distribution does not meet the above value, thus there is not enough evidence to conclude that the data have the tested distribution")
    st.write("Alpha means that the data have alpha percent correlation with the distribution")
    doc.add_paragraph("The minimum value of p value for alpha = 0.05 is 0.05")
    doc.add_paragraph("If the tested distribution does not meet the above value, thus there is not enough evidence to conclude that the data have the tested distribution")
    doc.add_paragraph("Alpha means that the data have alpha percent correlation with the distribution")
    # Perform distribution tests and display the results
    W, pvalNormal = normal_distribution_test(sorted_data)
    st.success("p-value for normal distribution test:"+str(round(pvalNormal, 2)))
    doc.add_paragraph("p-value for normal distribution test:"+str(round(pvalNormal, 2)))
    if pvalNormal > 0.05:
        st.info("Tested distribution exceeds 0.05, there is enough evidence")
        doc.add_paragraph("Tested distribution exceeds 0.05, there is enough evidence")
    else:
        st.warning("Tested distribution below 0.05, there is not enough evidence")
        doc.add_paragraph("Tested distribution below 0.05, there is not enough evidence")            
    W, pvalLogNormal = normal_distribution_test(log_data)
    st.success("p-value for log-normal distribution test:"+ str(round(pvalLogNormal, 2)))
    doc.add_paragraph("p-value for log-normal distribution test:"+ str(round(pvalLogNormal, 2)))
    B = exponential_test(sorted_data, len(sorted_data))
    st.success("B test value for exponential test:"+ str(round(B, 2)))
    doc.add_paragraph("B test value for exponential test:"+ str(round(B, 2)))
    M, dof1, dof2 = weibull_test(sorted_data, len(sorted_data))
    st.success("M value for Weibull distribution test:"+ str(round(M, 2)))
    doc.add_paragraph("M value for Weibull distribution test:"+ str(round(M, 2)))                
    t = np.arange(10001)
    hr, R, PDF, MTTF, RMTTF = cases(data, distribution, t)
    st.success("Predicted failure in operation hour:"+ str(round(MTTF, 2)))
    df['MTTF'] = MTTF
    doc.add_paragraph("Predicted failure in operation hour:"+ str(round(MTTF, 2)))
    # Display the plots
    fig = plot_results(hr, R, PDF, MTTF, RMTTF,data,distribution)
    st.subheader("Plots")
    st.pyplot(fig)
    fig.savefig("plot.png")
    doc.add_picture("plot.png",width=Inches(5))
    return df

def remove_text_between_parentheses(text):
    if isinstance(text, str):
        pattern = r"\([^()]*\)"
        return re.sub(pattern, "", text)
    return text

def remove_whitespace(sentence):
    if isinstance(sentence, str):
        sentence = re.sub(r'^\s+|\s+$', '', sentence)  # Remove leading/trailing whitespace
        sentence = re.sub(r'\s+', ' ', sentence)  # Replace multiple spaces with a single space
    return sentence

def process_all_values(df, column):
    if not isinstance(df, pd.DataFrame):
        raise ValueError("The input should be a DataFrame.")
    unique_values = df[column].unique()
    new_rows = []
    for index, row in df.iterrows():
        if row[column] == "all":
            for value in unique_values:
                new_row = row.copy()
                new_row[column] = value
                new_rows.append(new_row)
        else:
            new_rows.append(row)
    df_new = pd.DataFrame(new_rows)
    df_new = df_new[df_new[column] != "all"]
    return df_new

def process_df(df, column):
    if not isinstance(df, pd.DataFrame):
        raise ValueError("The input should be a DataFrame.")
    unique_values = df[column].unique()
    new_rows = []
    for index, row in df.iterrows():
        if str(row[column]).lower() == "all":
            for value in unique_values:
                new_row = row.copy()
                new_row[column] = value
                new_rows.append(new_row)
        else:
            new_rows.append(row)
    df_new = pd.DataFrame(new_rows)
    df_new = df_new.fillna(df_new.loc[df_new[column].str.lower() == "all"].iloc[0])
    df_new = df_new[df_new[column].str.lower() != "all"]
    df_new = df_new.reset_index(drop=True)
    return df_new

def mtbf_clc(doc):
    doc.add_heading("MTBF Calculation",level=1)
    choose=st.radio("How are you going to define unique componeny id?",('MPG SPG SSPG (BS EN 15380-2) + Train NUmber + TS + Cluster','Train Number + TS + Cluster','Cluster'),key=2)
    uploaded_file = st.file_uploader("Upload Filled delivery dates file, make sure you are using MM/DD/YYYY date format", type=["csv","xlsx"])
    updated_file = st.file_uploader("Upload Updated Cluster csv file", type=["csv","xlsx"])
    # Allow the user to upload the filled Excel sheet
    if uploaded_file and updated_file is not None:
        file_extension = uploaded_file.name.split('.')[-1]
        if file_extension.lower() == "xlsx":
            uploaded_df = pd.read_excel(uploaded_file)
        elif file_extension.lower() == "csv":
            uploaded_df = pd.read_csv(uploaded_file)
        else:
            st.error("Unsupported file format. Please upload either a CSV or XLSX file.")
        file_extension = updated_file.name.split('.')[-1]
        if file_extension.lower() == "xlsx":
            updated_df = pd.read_excel(updated_file)
        elif file_extension.lower() == "csv":
            updated_df = pd.read_csv(updated_file)
        else:
            st.error("Unsupported file format. Please upload either a CSV or XLSX file.")
        distribution = st.selectbox('Select Distribution Test?',('Normal', 'Log-Normal', 'Weibull'))
        doc.add_heading("Using "+str(distribution)+" distribution")
        daily_hours=st.number_input("Insert daily operating hours", min_value=0, max_value=24, value=20,step=1)
        doc.add_heading("Assuming it is operated "+str(daily_hours)+" hours per day")
        if st.button(f'Process MTBF Calculation using {distribution} Distribution'):                 
            uploaded_df[['Kereta', 'TS']] = uploaded_df['Trainset'].str.split('-', expand=True)
            uploaded_df = uploaded_df.drop(columns=['Trainset'])
            uploaded_df = uploaded_df.dropna(how='any')
            uploaded_df['Kereta'] = uploaded_df['Kereta'].str.strip()
            uploaded_df['TS'] = uploaded_df['TS'].str.strip()
            updated_df['TS'] = updated_df['TS'].astype(str).str.split('.').str[0]
            updated_df['Tanggal']=pd.to_datetime(updated_df['Tanggal']).dt.date
            # Find delivery time based on matching 'TS' and 'Kereta'
            merged_df = updated_df
            merged_df['Delivery Date'] = merged_df.apply(lambda row: 
                uploaded_df.loc[(uploaded_df['TS'] == row['TS']) & (uploaded_df['Kereta'] == row['Kereta']), 'Delivery Date'].values[0] 
                if row['TS'] in uploaded_df['TS'].values and row['Kereta'] in uploaded_df['Kereta'].values and uploaded_df.loc[(uploaded_df['TS'] == row['TS']) & (uploaded_df['Kereta'] == row['Kereta'])].shape[0] > 0 
                else '', axis=1)
            merged_df['Delivery Date'] = merged_df.apply(lambda row: uploaded_df.loc[(uploaded_df['TS'] == row['TS']) & (uploaded_df['Kereta'] == row['Kereta']), 'Delivery Date'].values[0] if row['TS'] in uploaded_df['TS'].values and row['Kereta'] in uploaded_df['Kereta'].values and uploaded_df.loc[(uploaded_df['TS'] == row['TS']) & (uploaded_df['Kereta'] == row['Kereta'])].shape[0] > 0 else '', axis=1)
            merged_df['Tanggal'] = pd.to_datetime(merged_df['Tanggal'])
            merged_df['Delivery Date'] = pd.to_datetime(merged_df['Delivery Date'])
            merged_df['Time Difference (days)'] = (merged_df['Tanggal'] - merged_df['Delivery Date']).dt.days
            merged_df['Tanggal'] = pd.to_datetime(merged_df['Tanggal']).dt.date
            merged_df['Delivery Date'] = pd.to_datetime(merged_df['Delivery Date']).dt.date
            merged_df['Time Difference (hours)']=merged_df['Time Difference (days)']*daily_hours
            if choose == "MPG SPG SSPG (BS EN 15380-2) + Train NUmber + TS + Cluster" and all(col in merged_df.columns for col in ['MPG', 'SPG', 'SSPG']):
                merged_df['component_id']=merged_df['Kereta']+'-'+merged_df['TS'].astype(str)+'-'+merged_df['MPG']+'-'+merged_df['SPG']+'-'+merged_df['SSPG']+'-'+merged_df['cluster_label']
            elif choose == "Train Number + TS + Cluster" and all(col in merged_df.columns for col in ['Kereta', 'TS' ]):
                merged_df['component_id']=merged_df['Kereta']+'-'+merged_df['TS'].astype(str)+'-'+merged_df['cluster_label']
            else:
                merged_df['component_id']=merged_df['cluster_label']           
            cols=['component_id']+merged_df.columns[:-1].tolist()
            merged_df=merged_df[cols]        
            merged_df['Time Difference (hours)'].dropna(inplace=True)
            merged_df = merged_df[merged_df['Time Difference (hours)'] >= 0]
            merged_df['Time Difference (days)'].dropna(inplace=True)
            merged_df = merged_df[merged_df['Time Difference (days)'] >= 0]    
            st.write(merged_df)
            doc.add_paragraph("Data:")
            # Add the dataframe as a table
            table = doc.add_table(merged_df.shape[0] + 1, merged_df.shape[1])
            table.style = 'Table Grid'  # Apply table grid style
            # Add column names to the table
            for i, column_name in enumerate(merged_df.columns):
                table.cell(0, i).text = column_name
            # Add data to the table
            for i, row in enumerate(merged_df.itertuples()):
                for j, value in enumerate(row[1:]):
                    table.cell(i + 1, j).text = str(value)        
            unique_klas = merged_df['component_id'].unique()
            # Create DataFrames based on unique 'Klas' values
            dfs = {}
            summary_df=pd.DataFrame(columns=["Nama Komponen","cluster_label","MTTF"])
            for component_id in unique_klas:
                if choose == "MPG SPG SSPG (BS EN 15380-2) + Train NUmber + TS + Cluster" and all(col in merged_df.columns for col in ['MPG', 'SPG', 'SSPG']):
                     df_klas = merged_df[merged_df['component_id'] == component_id][['TS', 'Tanggal', 'Kereta', 'Klasifikasi Gangguan', 'Nama Komponen', 'MPG', 'SPG', 'SSPG', 'cluster_label', 'Jumlah','Delivery Date','Time Difference (hours)','Time Difference (days)']].sort_values('Tanggal')
                else: 
                    df_klas = merged_df[merged_df['component_id'] == component_id][['TS', 'Tanggal', 'Kereta', 'Klasifikasi Gangguan', 'Nama Komponen', 'cluster_label', 'Jumlah','Delivery Date','Time Difference (hours)','Time Difference (days)']].sort_values('Tanggal')
                # Calculate Time To Failure (hours)
                df_klas['Time To Failure (hours)'] = 0
                df_klas['Tanggal'] = pd.to_datetime(df_klas['Tanggal']).dt.date
                df_klas['Delivery Date'] = pd.to_datetime(df_klas['Delivery Date']).dt.date
                df_klas.loc[df_klas.index[0], 'Time To Failure (hours)'] = (df_klas.loc[df_klas.index[0], 'Tanggal'] - df_klas.loc[df_klas.index[0], 'Delivery Date']).total_seconds() / 3600
                for i in range(1, len(df_klas)):
                    df_klas.loc[df_klas.index[i], 'Time To Failure (hours)'] = (df_klas.loc[df_klas.index[i], 'Tanggal'] - df_klas.loc[df_klas.index[i-1], 'Tanggal']).total_seconds() / 3600
                # Convert 'Jumlah' and 'Time To Failure (hours)' columns to integer
                df_klas['Jumlah'] = df_klas['Jumlah'].astype(int)
                df_klas['Time To Failure (hours)'].dropna(inplace=True)
                df_klas = df_klas[df_klas['Time To Failure (hours)'] > 0]
                df_klas['Time To Failure (hours)'] = df_klas['Time To Failure (hours)'].astype(int)
                dfs[f'{component_id}_df'] = df_klas
            # Access the created DataFrames
            for component_id, df_klas in dfs.items():
                if choose == "MPG SPG SSPG (BS EN 15380-2) + Train NUmber + TS + Cluster" and all(col in merged_df.columns for col in ['MPG', 'SPG', 'SSPG']):
                    df_klas = df_klas[['TS', 'Tanggal', 'Kereta', 'Klasifikasi Gangguan', 'Nama Komponen', 'MPG', 'SPG', 'SSPG', 'cluster_label', 'Delivery Date', 'Time Difference (hours)', 'Time Difference (days)', 'Jumlah', 'Time To Failure (hours)']]
                else:
                    df_klas = df_klas[['TS', 'Tanggal', 'Kereta', 'Klasifikasi Gangguan', 'Nama Komponen','cluster_label', 'Delivery Date', 'Time Difference (hours)', 'Time Difference (days)', 'Jumlah', 'Time To Failure (hours)']]
                train_number = df_klas['Kereta'].values[0]
                ts = df_klas['TS'].values[0]
                if choose == "MPG SPG SSPG (BS EN 15380-2) + Train NUmber + TS + Cluster" and all(col in merged_df.columns for col in ['MPG', 'SPG', 'SSPG']):
                    component_info = str(df_klas['MPG'].tolist()[0]) + '-' + str(df_klas['SPG'].tolist()[0]) + '-' + str(df_klas['SSPG'].tolist()[0]) + '-' + str(df_klas['cluster_label'].tolist()[0])
                else:
                    component_info = str(df_klas['cluster_label'].tolist()[0])
                if len(df_klas)>=3:
                    # Formatting the output
                    Train = "Train Number: {} - TS {}".format(train_number, ts)
                    Compo = "Component ID: {}".format(component_info)
                    # Displaying the output
                    st.subheader(Train)
                    st.subheader(Compo)
                    st.write("Information: Enough data to run test (available data {})".format(len(df_klas)))
                    st.write(df_klas)
                    doc.add_heading(Train, level=1)
                    doc.add_heading(Compo, level=1)
                    doc.add_heading("Information: Enough data to run test (available data {})".format(len(df_klas)))
                    # Add the dataframe as a table
                    table = doc.add_table(df_klas.shape[0] + 1, df_klas.shape[1])
                    table.style = 'Table Grid'  # Apply table grid style
                    # Add column names to the table
                    for i, column_name in enumerate(df_klas.columns):
                        table.cell(0, i).text = column_name
                    # Add data to the table
                    for i, row in enumerate(df_klas.itertuples()):
                        for j, value in enumerate(row[1:]):
                            table.cell(i + 1, j).text = str(value)
                    test(df_klas, distribution,doc)
                    df_klas['MTTF'] = df_klas['MTTF'].astype(str)
                    common_columns = list(set(summary_df.columns) & set(df_klas.columns))
                    summary_df = pd.concat([summary_df[common_columns], df_klas[common_columns]], ignore_index=True)
                else:
                    # Formatting the output
                    Train = "Train Number: {} - TS {}".format(train_number, ts) 
                    Compo = "Component ID: {}.".format( component_info)
                    # Displaying the output
                    st.subheader(Train)
                    st.subheader(Compo)
                    st.write("Information: Not Enough data to run test (minimum number of data: 3, available data {})".format(len(df_klas)))
                    df_klas['MTTF'] = "Not enough data"
                    df_klas['MTTF'] = df_klas['MTTF'].astype(str)
                    common_columns = list(set(summary_df.columns) & set(df_klas.columns))
                    summary_df = pd.concat([summary_df[common_columns], df_klas[common_columns]], ignore_index=True)
                    st.write(df_klas)
                    doc.add_heading(Train, level=1)
                    doc.add_heading(Compo, level=1)
                    doc.add_paragraph("Information: Not Enough data to run test (minimum number of data: 3, available data {})".format(len(df_klas)))
                    # Add the dataframe as a table
                    table = doc.add_table(df_klas.shape[0] + 1, df_klas.shape[1])
                    table.style = 'Table Grid'  # Apply table grid style
                    # Add column names to the table
                    for i, column_name in enumerate(df_klas.columns):
                        table.cell(0, i).text = column_name
                    # Add data to the table
                    for i, row in enumerate(df_klas.itertuples()):
                        for j, value in enumerate(row[1:]):
                            table.cell(i + 1, j).text = str(value)
            summary_df = summary_df.drop_duplicates(subset="cluster_label")
            summary_df = summary_df.reindex(columns=["Nama Komponen","cluster_label","MTTF"])
            doc2=Document()		
            doc2.add_heading("Summary", level=1)
	    # Add the dataframe as a table
            table = doc2.add_table(summary_df.shape[0] + 1, summary_df.shape[1])
            table.style = 'Table Grid'  # Apply table grid style
            # Add column names to the table
            for i, column_name in enumerate(summary_df.columns):
                table.cell(0, i).text = column_name
            # Add data to the table
            for i, row in enumerate(summary_df.itertuples()):
                for j, value in enumerate(row[1:]):
                    table.cell(i + 1, j).text = str(value)
            doc2_bytes = io.BytesIO()
            doc2.save(doc2_bytes)
            doc2_bytes.seek(0)		
            st.write(summary_df)
            st.download_button(
                label="Download summary docx",
                data=doc2_bytes.read(),
                file_name='summary.docx')
            for element in doc.element.body:
                doc2.element.body.append(element)
            doc_bytes = io.BytesIO()
            doc2.save(doc_bytes)
            doc_bytes.seek(0)
            #-----TO EDIT CLUSTERED
            st.download_button(
		label="Download report docx",
		data=doc_bytes.read(),
		file_name='report.docx')
    
def MTBF():
    st.empty()
    #Create a word doc
    doc=Document()
    st.title("Clustering and MTBF Calculator App")
    # File upload section
    st.subheader("Upload your Gangguan XLSX file")
    file = st.file_uploader("Upload XLSX file", type=["xlsx"])
    if file is not None:
        # Read the Excel file
        df = pd.read_excel(file)
        #get the table
        if df.columns[0][:5]=="Unnam":
            first_nonempty_row = df.index[df.notnull().any(axis=1)][0]
            # Delete the empty rows
            df = df.iloc[first_nonempty_row:]
            # Reset the index
            df.columns = df.iloc[0] 
            df = df[1:].reset_index(drop=True)
        # Find the first non-empty column index
        first_nonempty_column = df.columns[df.notnull().any(axis=0)][0]
        # Delete the empty columns
        df = df.loc[:, first_nonempty_column:]
        #strip string in dataframe
        df = df.applymap(lambda x: x.strip() if isinstance(x, str) else x)
        st.markdown("---")
        st.subheader("Pick your data")
        num_data = st.slider('How many of the data you are going to be used in clustering?', 0, len(df), 50)
        df=df.head(num_data)
        st.write(df)
        choose=st.radio("How are you going to define unique componeny id?",('MPG SPG SSPG (BS EN 15380-2) + Train NUmber + TS + Cluster','Train Number + TS + Cluster','Cluster'),key=1)
        # Create a new DataFrame with unique values in 'Nama Komponen'
        if choose == "MPG SPG SSPG (BS EN 15380-2) + Train NUmber + TS + Cluster" and all(col in merged_df.columns for col in ['MPG', 'SPG', 'SSPG']):
            default=["TS", "Tanggal", "Kereta", "Klasifikasi Gangguan", "Nama Komponen", "Jumlah","MPG","SPG", "SSPG"]
            PG=["MPG","SPG", "SSPG"]
        else:
            default=["TS", "Tanggal", "Kereta", "Klasifikasi Gangguan", "Nama Komponen", "Jumlah"]
            PG=[]
        options=st.multiselect("Select required columns", options=df.columns, default=default)
        required_columns=options
        df = df[required_columns]
        df=df.drop_duplicates(subset=['Tanggal','TS','Kereta','Nama Komponen','Jumlah']+PG)
        unique_count = df['Nama Komponen'].nunique()  # Get the number of unique values
        st.write("Number of unique values in 'Nama Komponen':", unique_count)
        #create df of combinations Kereta and TS
        df_delivery = process_all_values(df, "Kereta")
        df_delivery = pd.DataFrame({
            'Trainset': df_delivery.apply(lambda row: f"{row['Kereta']} - {row['TS']}", axis=1),
            'Delivery Date': [''] * len(df_delivery)
        })
        df_delivery = df_delivery.drop_duplicates(subset=['Trainset'])
        df_delivery = df_delivery.rename(columns={'Trainset': 'Trainset', 'Delivery Date': 'Delivery Date'})
        delivery_file= df_delivery.to_csv(index=False).encode('utf-8')   
        dfx = process_df(df,"Kereta")
        dfx=dfx.drop_duplicates(subset=['Tanggal','TS','Kereta','Nama Komponen','Jumlah']+PG)
        if st.button('Cluster!'):
            dfx['original komponen name'] = dfx['Nama Komponen']
            # Apply text preprocessing
            dfx['Nama Komponen'] = dfx['Nama Komponen'].apply(remove_whitespace)
            dfx['Nama Komponen'] = dfx['Nama Komponen'].apply(remove_text_between_parentheses)
            dfx['Nama Komponen'] = dfx['Nama Komponen'].apply(lambda x: x.lower() if isinstance(x, str) else x)
            # Preprocess the text data
            vectorizer = TfidfVectorizer(stop_words='english')
            X = vectorizer.fit_transform(dfx['Nama Komponen'].values.astype('U'))
            max_clusters = int(round(0.8 * dfx['Nama Komponen'].nunique(), 0))
            best_score = -1
            best_num_clusters = 0
            for num_clusters in range(2, max_clusters + 1):
                kmeans = MiniBatchKMeans(n_clusters=num_clusters, random_state=0)
                cluster_labels = kmeans.fit_predict(X)
                silhouette_avg = silhouette_score(X, cluster_labels, sample_size=1000)
                if silhouette_avg > best_score:
                    best_score = silhouette_avg
                    best_num_clusters = num_clusters
            st.subheader("Clustering result")        
            st.write("Optimal number of clusters:", best_num_clusters)
            # Apply K-means clustering with the optimal number of clusters
            kmeans = KMeans(n_clusters=best_num_clusters, random_state=0)
            kmeans.fit(X)
            # Get the cluster labels
            labels = kmeans.labels_
            # Add the cluster labels as a column next to 'Nama Komponen'
            dfx['cluster_label'] = labels
            # Create a dictionary to store the first value in each cluster
            cluster_first_values = {}
            for cluster_label in set(labels):
                cluster_values = dfx.loc[dfx['cluster_label'] == cluster_label, 'Nama Komponen']
                first_value = cluster_values.iloc[0]
                cluster_first_values[cluster_label] = first_value
            # Update the 'cluster_label' column with the first value in each cluster
            dfx['cluster_label'] = dfx['cluster_label'].map(cluster_first_values)
            # Move the 'Jumlah' column to the rightmost position
            columns = dfx.columns.tolist()
            columns.remove('Jumlah')
            columns.append('Jumlah')
            dfx = dfx[columns]
            st.write(dfx)
            # Download button
            csv_data = dfx.to_csv(index=False).encode('utf-8')
            #-----TO EDIT CLUSTERED
            st.download_button(
                                label="We have clustered the component's name using Machine Learning, if you want to edit it, Click to download data as CSV then upload it again",
                                data=csv_data,
                                file_name='cluster.csv',
                                mime='text/csv',)
        st.markdown("---") 
                # Provide a download button for the Excel file
        st.subheader("Please click to download the delivery date csv file and fill in the delivery dates")
        st.download_button(
                            label="Download the delivery csv file template",
                            data=delivery_file,
                            file_name='delivery_template.csv',
                            mime='text/csv',)
        mtbf_clc(doc)
    else:
        st.markdown("---") 
        st.subheader("Or if you already have filled delivery data and cluster data, upload to the following")
        mtbf_clc(doc)
	
##-- CHAT	
def chat():
	st.empty()
	df = pd.read_pickle('./standards.pkl')
	df['num_chars'] = df['text'].apply(lambda x: len(x))
	df = df[df['num_chars'] != 0]
	# Choose a topic
	st.write("""This App uses AI, though sometimes it provides correct answer, sometimes it may not. Always use your own discretion.
		This AI only fit for a short question answering 
		This AI uses paid API, get your openai api key [here](https://platform.openai.com/account/api-keys)""")
	OPENAI_API_KEY=st.text_input("insert openai api",type="password")
	unique_values = set(df["location"].str.split("/").str[1])
	std_type = st.multiselect('Select Standards',unique_values,unique_values)
	keyword = st.text_input("choose topic","running dynamic")
	query = st.text_input("insert query","vehicle at what speed that must perform dynamic performance test?")
	if st.button("Process"):
	# Filter by keyword
		filtered_std  = df[df["location"].apply(lambda x: any(item in x for item in std_type))]
		filtered_std = filtered_std[filtered_std['text'].str.contains(keyword, flags=re.IGNORECASE)]
		selected_df = filtered_std[["location", "name", "id"]]
		selected_df['link'] = selected_df['id'].apply(lambda x: f'<a target="_blank" href="https://drive.google.com/file/d/{x}/view">{x}</a>')
		selected_df = selected_df.drop("id", axis=1)

		st.write(selected_df.shape[0])
		if selected_df.shape[0] > 0:
			selected_df = selected_df.to_html(escape=False)
			st.write(selected_df, unsafe_allow_html=True)
			joined = ",".join(filtered_std['text'].astype(str))
			doc = Document(page_content=joined)
			text_splitter = RecursiveCharacterTextSplitter(chunk_size = 1000,chunk_overlap  = 20,length_function = len)
			texts = text_splitter.split_documents([doc])
			embeddings = OpenAIEmbeddings(openai_api_key=OPENAI_API_KEY)
			docsearch = Chroma.from_documents(texts, embeddings)
			qa = RetrievalQA.from_chain_type(llm=OpenAI(openai_api_key=OPENAI_API_KEY), chain_type="map_rerank", retriever=docsearch.as_retriever(),return_source_documents=True)
			result = qa({"query": query})
			st.write("Answer :")
			st.write(result["result"])
			st.markdown("---")
			st.write("Sources :")
			source_documents = [doc.page_content for doc in result["source_documents"]]
			unique_sources = pd.concat([df[df["text"].str.contains(max(doc.split("."), key=len).strip(), case=False)]["location"] for doc in source_documents]).unique()
			locations_string = "\n".join(unique_sources)
			st.write(locations_string)
		else:
			st.write("No data contain specific keyword")

def req():
	st.empty()
	st.write("""This App uses AI, though sometimes it provides correct answer, sometimes it may not. Always use your own discretion.
		This AI will gather information from selected standards and provides requrements each component must comply
		This AI uses paid API, get your openai api key [here](https://platform.openai.com/account/api-keys)""")
	OPENAI_API_KEY=st.text_input("insert openai api",type="password")
	df = pd.read_pickle('./standards.pkl')
	df['num_chars'] = df['text'].apply(len)
	df = df[df['num_chars'] != 0]
	unique_values = set(df["location"].str.split("/").str[1])
	std_type = st.multiselect('Select Standards',unique_values,unique_values)
	component = st.text_input("insert component's name","Vehicle body")
	from langchain.prompts import PromptTemplate, StringPromptTemplate
	template = """
		{summaries}
		{question}
		"""
	question = """
		You are a quality control engineer responsible for ensuring compliance with industry standards for a {component}. 
		Your task is to develop a set of parameters that all instances of the {component} must meet in order to comply with the given standards.
		Write a detailed description of {component} and the specific standards that apply to it. 
		Outline the key parameters that must be considered and provide a clear explanation of how each parameter contributes to compliance.
		Your response:
		"""
	prompt = PromptTemplate.from_template(template)
	question=prompt.format(component=component)
	if st.button("Process"):
		filtered_std  = df[df["location"].apply(lambda x: any(item in x for item in std_type))]
		filtered_std = filtered_std[filtered_std['text'].str.contains(component, flags=re.IGNORECASE)]
		selected_df = filtered_std[["location", "name", "id"]]
		selected_df['link'] = selected_df['id'].apply(lambda x: f'<a target="_blank" href="https://drive.google.com/file/d/{x}/view">{x}</a>')
		selected_df = selected_df.drop("id", axis=1)
		
		st.write(selected_df.shape[0])
		if selected_df.shape[0]>0:
			selected_df=selected_df.head(5)
			selected_df = selected_df.to_html(escape=False)
			st.write(selected_df, unsafe_allow_html=True)
			joined = ",".join(filtered_std['text'].astype(str))
			doc = Document(page_content=joined)
			text_splitter = RecursiveCharacterTextSplitter(chunk_size = 1000,chunk_overlap  = 20,length_function = len)
			texts = text_splitter.split_documents([doc])
			embeddings = OpenAIEmbeddings(openai_api_key=OPENAI_API_KEY)
			docsearch = Chroma.from_documents(texts, embeddings)
			from langchain.chat_models import ChatOpenAI
			from langchain.chains import LLMChain
			qa = RetrievalQAWithSourcesChain.from_chain_type(llm=OpenAI(openai_api_key=OPENAI_API_KEY), chain_type="stuff", retriever=docsearch.as_retriever(),return_source_documents=True,chain_type_kwargs={"prompt": PromptTemplate(template=template,input_variables=["summaries", "question"])})
			result = qa.run(component_name)
			st.write("Answer :")
			st.write(result["result"])
			st.markdown("---")
			st.write("Sources :")
			source_documents = [doc.page_content for doc in result["source_documents"]]
			unique_sources = pd.concat([df[df["text"].str.contains(max(doc.split("."), key=len).strip(), case=False)]["location"] for doc in source_documents]).unique()
			locations_string = "\n".join(unique_sources)
			st.write(locations_string)
		else:
			st.write("No standards contain specific keyword")
st.empty()		
page_names_to_funcs = {
    "Product Breakdown Structure": system_requirement,
    "Material Code":Matcod,
    "Initial FMECA":FMECA,
    "Failure Rate Calculator":failure,
    "Function Breakdown Structure":FBS,
    "Standards finder":Standards,
    "Possible Supplier":Supplier,
    "Component Clustering & MTBF Calculator":MTBF,
    "Talk To Your Standards":chat,
    #"Requirements for each component":req
    }

selected_page = st.sidebar.radio("Select a page", page_names_to_funcs.keys())
page_names_to_funcs[selected_page]()
